from PyQt5 import QtWidgets, QtCore, QtGui, QtPrintSupport
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.uic import *
from docx import *
from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.section import WD_ORIENT, WD_SECTION

import os
from os import path
import sys
from datetime import *
import time
import sqlite3
import random
import multiprocessing as mp
import os
from stat import S_IREAD, S_IRGRP, S_IROTH
from fpdf import FPDF


today = date.today()

print("Today's date:", today)


first_open_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "first_open_win.ui"))
home_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "home.ui"))
splash_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "splash.ui"))
virification_alert_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "virification_alert.ui"))
setting_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "setting.ui"))
history_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "history.ui"))
add_new_pea_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "add_new_pea.ui"))
add_new_art_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "add_new_art.ui"))
delete_pea_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "delet_pea.ui"))
fix_kridi_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "fix_kridi.ui"))
fix_kridi_history_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "fix_history.ui"))
kridi_history_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "kridi_history.ui"))
edite_pea_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "edit_pea.ui"))
edite_art_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "edit_art.ui"))
#TODO : # help_win_dir,_ = loadUiType(path.join(path.dirname(__file__), "help.ui"))


conn = sqlite3.connect('src/db.db')
curs = conn.cursor()
curs.execute('CREATE TABLE IF NOT EXISTS user (NAME TEXT, PASSWORD TEXT)')

header_ = 'this is a header'
footer_ = 'this is a footer'

class Splash(QWidget, splash_win_dir):#DONE
    def __init__(self, parent = None):
        super(Splash, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('مرحبا')
        self.timer = QBasicTimer()
        self.step = 0
        self.prog()
    def prog(self):
        if self.timer.isActive():
            self.timer.stop()
        else:
            self.timer.start(100, self)
    def timerEvent(self, event):
        if self.step >= 100 :
            self.timer.stop()
            curs.execute('SELECT first FROM tools')
            first_stat = curs.fetchone()
            if first_stat[0] == 1:
                print('its the first open ')
                wilco_msg = QMessageBox.information(self, '', "بما أنها المرة الأولى هناك بعض المعلومات يجب تحديثها ", QMessageBox.Ok)
                if wilco_msg == QMessageBox.Ok:
                    self.first_open_w = FirstOpen()
                    self.first_open_w.show()
                    self.close()

            else:
                curs.execute('SELECT safe FROM tools ')
                curent_safe_stat = curs.fetchone()
                if curent_safe_stat[0] == 0:
                    print('its not the the first open ')
                    self.virification_alert_wn = VirificationAlert()
                    self.virification_alert_wn.show()
                    self.close()
                else:
                    self.home_wn = Home()
                    self.home_wn.show()
                    self.close()
           
            return
        self.step += 4
        self.progressBar.setValue(self.step)

class FirstOpen(QMainWindow, first_open_win_dir):# done
    def __init__(self, parent = None):
        super(FirstOpen, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('نقطة البداية')
        self.save_btn.clicked.connect(self.save_info_first_open)
        self.pass_entry.setEchoMode(QLineEdit.Password)
        self.pass_confirmation.setEchoMode(QLineEdit.Password)


    def save_info_first_open(self):
        if self.user_name_entry.text() == '' or self.user_name_entry.text() == ' ' or self.user_name_entry.text() == 0 :
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('لايمكن ترك إسم المستخدم فارغ ')
            self.err.setWindowTitle('خطأ')
            self.user_name_entry.setFocus()

        elif self.safe_mod_checkBox.isChecked():
            curs.execute('INSERT INTO user (NAME, PASSWORD) VALUES ( "{}", "{}" ); '.format(self.user_name_entry.text(), ''))
            curs.execute('UPDATE tools SET user = "{}" , safe = 1  ;'.format(self.user_name_entry.text()))
            curs.execute('UPDATE tools SET first = {}'.format(2))
            conn.commit()
            wilcome_msg = QMessageBox.information(self, '', "تم تحديث المعلومات بنجاح", QMessageBox.Ok)
            if wilcome_msg == QMessageBox.Ok:
                self.home_wn = Home()
                self.home_wn.show()
                self.close()
        else:
            if self.pass_entry.text() == '' or self.pass_entry.text() == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('كلمة مرور غير صالحة ')
                self.err.setWindowTitle('خطأ')
                self.pass_entry.setFocus()

            elif self.pass_confirmation.text() == '' or self.pass_confirmation.text() == ' ' or self.pass_confirmation.text() != self.pass_entry.text():
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('الكلمات غير متطابقة ')
                self.err.setWindowTitle('خطأ')
                self.pass_confirmation.setFocus()
            else:
                curs.execute('INSERT INTO user (NAME, PASSWORD) VALUES ( "{}", "{}" ) '.format(self.user_name_entry.text(),self.pass_entry.text()))
                curs.execute('UPDATE tools SET user = "{}", first = {} , safe = 0 '.format(self.user_name_entry.text(), 2))
                conn.commit()
                self.user_name_entry.setText('')
                self.pass_entry.setText('')
                self.pass_confirmation.setText('')
                wilcome_msg = QMessageBox.information(self, '', "تم تحديث المعلومات بنجاح", QMessageBox.Ok)
                if wilcome_msg == QMessageBox.Ok:
                    self.home_wn = Home()
                    self.home_wn.show()
                    self.close()


class Print_pdf:
    def __init__(self, file_name, table_head, table = '', spacing = 1):
        ypdf = FPDF()
        ypdf.add_font('DejaVu', '', 'src/DejaVuSansCondensed.ttf', uni=True)
        ypdf.set_font('DejaVu', '', 11)
        # ypdf.set_font("Arial", size=10)
        ypdf.add_page()
        # #  # Add an address
        ypdf.cell(100)
        ypdf.cell(0, 5, 'Mike Driscoll', ln=1)
        ypdf.cell(100)
        ypdf.cell(0, 5, '123 American Way', ln=1)
        ypdf.cell(100)
        ypdf.cell(0, 5, 'Any Town, USA', ln=1)

        # #     # # Line break
        ypdf.ln(20)
        curs.execute('SELECT * FROM "{}"'.format(table))
        data = curs.fetchall()

        data.insert(0, table_head)

        col_width = ypdf.w / (len(table_head) + 0.5)
        row_height = ypdf.font_size
        for row in data:
            for item in row:
                ypdf.cell(col_width, row_height * spacing, txt=str(item), border=1, align='C')
            ypdf.ln(row_height * spacing)


        # !TODO fix this

        if len(file_name.split('/')[-1].split('.')) == 2:
            if str(file_name.split('/')[-1].split('.')[1]) == 'pdf' :
                print('file name : ', file_name)
                ypdf.output(file_name, 'F')
                # self.mk_pdf(self.file_name)
                # pdf.output('/mnt/AC72F2C272F29076/works/stock/stock-management-V2/src/uuuuuuuuuu.pdf')

            else:
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('صيغة الملف غير مقبولة')
                self.err.setWindowTitle('خطأ')

        else:
            print(' no extention')
            file_ = str(file_name) + '.pdf'
            print('file name : ', file_)
            ypdf.output(file_, 'F')

class Selles_history(QWidget, history_win_dir):#DONE
    def __init__(self, parent = None):
        super(Selles_history, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('سجل المبيعات')
        self.refresh_()
        self.back_from_history.clicked.connect(self.back_home)
        self.print_history.clicked.connect(self.pprint)



    #? for word 
#     """
#     def pprint(self, location):#TODO PRINT data 
#         def fill_doc( file_name, data, h = '', f = ''):
#             from docx.shared import Inches
#             doc = Document()
#             # doc.add_heading('test heading')
#             section = doc.sections[0]
#             heade = section.header
#             paragraph = heade.paragraphs[0]
#             paragraph.text = h
#             footer = section.footer
#             footer.paragraphs[0].text = f
#             # for i in curs.fetchall():
#             #     ii.appe   nd(i)
#             # add table ------------------
#             curs.execute('SELECT ID, date, name, article, qt, total FROM selles_history ORDER BY ID DESC')
#             statu_ = curs.fetchall()
#             table = doc.add_table(1, 6)
#             # current_section = doc.sections[-1]
#             # new_width, new_height = current_section.page_height, current_section.page_width
#             # new_section = doc.add_section(WD_SECTION.NEW_PAGE)
#             # new_section.orientation = WD_ORIENT.LANDSCAPE
#             # new_section.page_width = new_width
#             # new_section.page_height = new_height
#             # populate header row --------
#             heading_cells = table.rows[0].cells
#             header_titles =['رقم الزبون','تاريخ العملية', 'الإسم', 'السلعة', 'الكمية', 'المجموع']
#             # heading_cells[0].text = 'ID'
#             # heading_cells[1].text = 'date'
#             # heading_cells[2].text = 'name'
#             # heading_cells[3].text = 'article'
#             # heading_cells[4].text = 'qt'
#             # heading_cells[5].text = 'total'
#             # heading_cells[6].text = 'pay_date'
#             for i in range(6):
#                 heading_cells[i].text = header_titles[i]

#             # add a data row for each item
#             for item in data:
#                 cells = table.add_row().cells
#                 cells[0].text = str(item[0])
#                 cells[1].text = str(item[1])
#                 cells[2].text = str(item[2])
#                 cells[3].text = str(item[3])
#                 cells[4].text = str(item[4])
#                 cells[5].text = str(item[5]) + ' DH'
#                 # cells[6].text = str(item[6])
    
#             table.style = 'LightShading-Accent1'
#             # fff =file_name + '.docx'
#             #TODO 
#             doc.save(file_name)
#             print(file_name + ' : saved successfully')
#             os.chmod(file_name, S_IREAD)
#             print(file_name + ' : changed to read only')
#             # os.system(fff)
#             # print(fff.split('/')[-1])
#         file_name,_ = QFileDialog.getSaveFileName(self, caption = 'حفظ في :', directory = '.', filter = "text files (*.doc *.docx)")
#         if file_name :
#             curs.execute('SELECT * FROM selles_history')
#             ii = curs.fetchall()
#             print(file_name)
        
#             # print(len(file_name.split('/')[-1].split('.')))
    
#             if len(file_name.split('/')[-1].split('.')) == 2:
#                 if file_name.split('/')[-1].split('.')[1] == 'doc' or  file_name.split('/')[-1].split('.')[1] == 'docx':
#                     # print('extention : doc or docx')#TODO header & footer
#                     fill_doc(file_name, ii, header_, footer_)
#                     QApplication.processEvents()
#                 else:
#                     self.err = QtWidgets.QErrorMessage()
#                     self.err.showMessage('صيغة الملف غير مقبولة')
#                     self.err.setWindowTitle('خطأ')
              
#             else:
#                 print(' no extention')
#                 file_ = file_name + '.docx'
#                 fill_doc(file_, ii, header_, footer_)
#                 QApplication.processEvents()
#         else:
#             print('printing canceled')
# """
#? for pdf 

    def pprint(self, spacing = 1):#TODO PRINT data
        file_name, _ = QFileDialog.getSaveFileName(self, caption='حفظ في :', directory='.', filter="text files (*.pdf)")
        if file_name:
            print_pdf = Print_pdf(file_name, ['ID', 'date', 'name', 'art', 'qt', 'total', 'pay_date', 'test'], 'selles_history')


        else:
            print('printing canceled')

    def back_home(self):
        self.close()
        self.home_wn = Home()
        self.home_wn.show()

    def refresh_(self):
        while self.selles_history_table.rowCount() > 0 :
            self.selles_history_table.removeRow(0)
        curs.execute('SELECT  date, client, article, qt, payed, not_payed_yet, pay_date FROM selles_history ORDER BY ID DESC')
        rus_ = curs.fetchall()
        self.selles_history_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.selles_history_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.selles_history_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

class Buyes_history(QWidget, history_win_dir):#DONE
    def __init__(self, parent = None):
        super(Buyes_history, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('سجل المشتريات')
        self.refresh_()
        self.back_from_history.clicked.connect(self.back_home)
        self.print_history.clicked.connect(self.pprint)


#TODO : here we are again
    def pprint(self):#TODO PRINT data 
        pass


    def back_home(self):
        self.close()
        self.home_wn = Home()
        self.home_wn.show()

    def refresh_(self):
        while self.selles_history_table.rowCount() > 0 :
            self.selles_history_table.removeRow(0)
        curs.execute('SELECT  date, seller, article, qt, payed, not_payed_yet, pay_date FROM buys_history ORDER BY ID DESC')
        rus_ = curs.fetchall()
        self.selles_history_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.selles_history_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.selles_history_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        print(curs.fetchall())

class ADD_new_client(QWidget, add_new_pea_win_dir):#DONE
    def __init__(self, parent = None):
        super(ADD_new_client, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('إظافة زبون')
        self.cancel_btn.clicked.connect(self.cancel)
        self.save_new_client.clicked.connect(self.save_new_client_)

        self.clients_full_names = []
    def save_new_client_(self):
        if self.F_name.text() == '' or self.F_name.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الإسم الأول ')
            self.err.setWindowTitle('خطأ')
            self.F_name.setFocus()

        elif self.L_name.text() == '' or self.L_name.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الإسم الأخير ')
            self.err.setWindowTitle('خطأ')
            self.L_name.setFocus()

        else:
            self.clients_cne_list = []
            curs.execute('SELECT cne FROM clients')
            for cne in curs.fetchall():
                self.clients_cne_list.append(cne[0])

            curs.execute('SELECT F_name, L_name FROM clients')
            for fname, lname in curs.fetchall():
                self.clients_full_names.append(fname + ' ' + lname)

            if self.F_name.text() + ' ' + self.L_name.text() in self.clients_full_names:
                current_info = curs.execute('SELECT F_name, L_name, cne FROM clients'
                                            ' WHERE F_name LIKE "{}" AND L_name LIKE "{}";'.format(self.F_name.text(),
                                                                                                   self.L_name.text()))
                self.current = []
                for oo in current_info.fetchall():
                    print('00 = {}'.format(oo))
                    self.current.append(oo[0])
                    self.current.append(oo[1])
                    self.current.append(oo[2])
                
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage(' هذا الزبون موجود بالفعل في قاعدة البيانات \n #الإسم : " ‏؜{} "  #اللقب : " {} "  #البطاقة الوطنية : " {} " '
                    .format(str(self.current[0]), str(self.current[1]), str(self.current[2])))
                self.err.setWindowTitle('هذا الزبون موجود في قاعدة البيانات')
                self.F_name.setFocus()

            elif self.cne.text() in self.clients_cne_list:
                bb = curs.execute('SELECT F_name, L_name FROM clients WHERE cne LIKE "{}"'.format(self.cne.text()))
                liss = bb.fetchone()
                msg = 'هذه البطاقة الوطنية هي للسيد : {} {}'
                self.err = QtWidgets.QErrorMessage()
                self.err.setWindowTitle('هناك تداخل في البيانات')
                self.err.showMessage(msg.format(liss[0], liss[1]))

            else:
                if self.cne.text() == '' or self.cne.text() == ' ':
                    self.cne.setText('-')

                if self.note.text() == '' or self.note.text() == ' ':
                    self.note.setText('-')
                
                curs.execute('INSERT INTO clients (F_name, L_name, cne, note) VALUES("{}", "{}", "{}", "{}");'
                             .format(self.F_name.text(), self.L_name.text(), self.cne.text(), self.note.text()))

                conn.commit()
                self.F_name.setText('')
                self.F_name.setFocus()
                self.L_name.setText('')
                self.cne.setText('')
                self.note.setText('')

    def cancel(self):
        self.close()
        self.home_wn = Home()
        self.home_wn.show()

class ADD_new_seller(QWidget, add_new_pea_win_dir):# DONE
    def __init__(self, parent = None):
        super(ADD_new_seller, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('إظافة موزع')
        self.cancel_btn.clicked.connect(self.cancel)
        self.save_new_client.clicked.connect(self.save_new_client_)

        self.sellers_full_names = []
    def save_new_client_(self):
        if self.F_name.text() == '' or self.F_name.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الإسم الأول ')
            self.err.setWindowTitle('خطأ')
            self.F_name.setFocus()

        elif self.L_name.text() == '' or self.L_name.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الإسم الأخير ')
            self.err.setWindowTitle('خطأ')
            self.L_name.setFocus()

        else:
            self.sellers_cne_list = []
            curs.execute('SELECT cne FROM sellers')
            for cne in curs.fetchall():
                self.sellers_cne_list.append(cne[0])

            curs.execute('SELECT F_name, L_name FROM sellers')
            for fname, lname in curs.fetchall():
                self.sellers_full_names.append(fname + ' ' + lname)

            if self.F_name.text() + ' ' + self.L_name.text() in self.sellers_full_names:
                current_info = curs.execute('SELECT F_name, L_name, cne FROM sellers WHERE F_name LIKE "{}" AND L_name LIKE "{}";'.format(self.F_name.text(), self.L_name.text()))
                self.current = []
                for oo in current_info.fetchall():
                    self.current.append(oo[0])
                    self.current.append(oo[1])
                    self.current.append(oo[2])
                    
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage(
                    ' هذا الزبون موجود بالفعل في قاعدة البيانات \n #الإسم : " ‏؜{} "  #اللقب : " {} "  #البطاقة الوطنية : " {} " '
                    .format(str(self.current[0]), str(self.current[1]), str(self.current[2])))
                self.err.setWindowTitle('هذا الزبون موجود في قاعدة البيانات')
                self.F_name.setFocus()

            elif self.cne.text() in self.sellers_cne_list:
                bb = curs.execute('SELECT F_name, L_name FROM sellers WHERE cne LIKE "{}"'.format(self.cne.text()))
                liss = bb.fetchone()
                msg = 'هذه البطاقة الوطنية هي للسيد : {} {}'
                self.err = QtWidgets.QErrorMessage()
                self.err.setWindowTitle('هناك تداخل في البيانات')
                self.err.showMessage(msg.format(liss[0], liss[1]))

            else:
                if self.cne.text() == '' or self.cne.text() == ' ':
                    self.cne.setText('-')

                if self.note.text() == '' or self.note.text() == ' ':
                    self.note.setText('-')
                
                curs.execute('INSERT INTO sellers (F_name, L_name, cne, note) VALUES("{}", "{}", "{}", "{}");'
                             .format(self.F_name.text(), self.L_name.text(), self.cne.text(), self.note.text()))

                conn.commit()
                self.F_name.setText('')
                self.F_name.setFocus()
                self.L_name.setText('')
                self.cne.setText('')
                self.note.setText('')

    def cancel(self):
        self.close()
        self.home_wn = Home()
        self.home_wn.show()

class ADD_new_art(QWidget, add_new_art_win_dir):# DONE
    def __init__(self, parent = None):
        super(ADD_new_art, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('إظافة منتوج')
        self.cancel_btn.clicked.connect(self.cancel)
        self.save_new_art.clicked.connect(self.save_new_art_)

        
        self.art_list = []
    def save_new_art_(self):
        if self.name.text() == '' or self.name.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الإسم  ')
            self.err.setWindowTitle('خطأ')
            self.name.setFocus()

        elif self.qt.text() == '' or self.qt.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الكمية ')
            self.err.setWindowTitle('خطأ')
            self.qt.setFocus()

        elif self.price.text() == '' or self.price.text() == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال الثمن ')
            self.err.setWindowTitle('خطأ')
            self.price.setFocus()


        else:
            self.current = []
            curs.execute('SELECT name FROM articles')
            for nn in curs.fetchall():
                self.art_list.append(nn[0])

            if self.name.text() in self.art_list:
                curs.execute('SELECT name, qt, price FROM articles WHERE name LIKE "{}" '.format(self.name.text()))
                
                for oo in curs.fetchone():
                    self.current.append(oo)
                    
                print(self.current)
                
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('" {} " هي موجودة بالفعل في قاعدة البيانات و كميتها الحالية : "{}" و ثمنها : {} درهم'
                    .format(self.current[0], self.current[1], self.current[2]))
                self.err.setWindowTitle('هذه السلعة موجود في قاعدة البيانات')
                self.name.setFocus()
                
            else:
                if self.note.text() == '' or self.note.text() == ' ':
                    self.note.setText('-')
                if self.type.text() == '' or self.type.text() == ' ':
                    self.type.setText('-')

                curs.execute('INSERT INTO articles (name, type, qt, price, note) VALUES("{}", "{}", {}, {}, "{}");'
                             .format(self.name.text(), self.type.text(), self.qt.text(), self.price.text(), self.note.text()))

                conn.commit()
                self.name.setText('')
                self.name.setFocus()
                self.type.setText('')
                self.qt.setText('')
                self.price.setText('')
                self.note.setText('')

    def cancel(self):
        self.close()
        self.home_wn = Home()
        self.home_wn.show()

class Remove_client(QWidget, delete_pea_win_dir): # DONE
    def __init__(self, parent = None):
        super(Remove_client, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('حذف زبون')
        self.fill_combo()
        self.shose_pea_to_rmv.currentTextChanged.connect(self.set_cne)
        self.rmv_btn.clicked.connect(self.delete)
        self.cancel_rmv.clicked.connect(self.cncl)



    def set_cne(self):
        if self.shose_pea_to_rmv.currentText() != '':
            curs.execute('SELECT cne FROM clients WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
                         .format(self.shose_pea_to_rmv.currentText().split('|')[0], self.shose_pea_to_rmv.currentText().split('|')[1]))
            
            
            self.show_cne_rmv_page.setText(curs.fetchone()[0])

    def fill_combo(self):
        clients_names_list_ = ['']
        curs.execute('SELECT F_name, L_name  FROM clients')
        for fn, ln in curs.fetchall():
            clients_names_list_.append(fn + '|' + ln)
            
        self.shose_pea_to_rmv.addItems(clients_names_list_)

    def delete(self):
        
        curs.execute('DELETE FROM clients WHERE F_name LIKE "{}" AND L_name LIKE "{}" AND cne LIKE "{}" '
                     .format(self.shose_pea_to_rmv.currentText().split('|')[0],
                             self.shose_pea_to_rmv.currentText().split('|')[1],
                     self.show_cne_rmv_page.text()))
        self.shose_pea_to_rmv.clear()
        self.fill_combo()
        self.shose_pea_to_rmv.setCurrentIndex(0)
        self.show_cne_rmv_page.setText('')
        conn.commit()




    def cncl(self):
        self.close()
        self.home = Home()
        self.home.show()

class Remove_seller(QWidget, delete_pea_win_dir): # DONE
    def __init__(self, parent = None):
        super(Remove_seller, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('حذف موزع')
        self.fill_combo()
        self.shose_pea_to_rmv.currentTextChanged.connect(self.set_cne)
        self.rmv_btn.clicked.connect(self.delete_)
        self.cancel_rmv.clicked.connect(self.cncl)

    def set_cne(self):
        if self.shose_pea_to_rmv.currentText() != '':
            curs.execute('SELECT cne FROM sellers WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
                         .format(self.shose_pea_to_rmv.currentText().split('|')[0], self.shose_pea_to_rmv.currentText().split('|')[1]))

            self.show_cne_rmv_page.setText(curs.fetchone()[0])
    def fill_combo(self):
        clients_names_list_ = ['']
        curs.execute('SELECT F_name, L_name  FROM sellers')
        for fn, ln in curs.fetchall():
            clients_names_list_.append(fn + '|' + ln)
            
        self.shose_pea_to_rmv.addItems(clients_names_list_)

    def delete_(self):
        curs.execute('DELETE FROM sellers WHERE F_name LIKE "{}" AND L_name LIKE "{}" AND cne LIKE "{}" '
                     .format(self.shose_pea_to_rmv.currentText().split('|')[0],
                             self.shose_pea_to_rmv.currentText().split('|')[1],
                     self.show_cne_rmv_page.text()))
        self.shose_pea_to_rmv.clear()
        self.fill_combo()
        self.shose_pea_to_rmv.setCurrentIndex(0)
        self.show_cne_rmv_page.setText('')
        conn.commit()

    def cncl(self):
        self.close()
        self.home = Home()
        self.home.show()

class C_kridi_fix_history(QWidget, fix_kridi_history_win_dir):
    def __init__(self, parent = None):
        super(C_kridi_fix_history, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('سجل القروض المدفوعة')
        self.refresh_()
        self.back_btn.clicked.connect(self.back__)
        self.print_btn.clicked.connect(self.pprint)



    def pprint(self):#TODO PRINT data 
        pass



    def back__(self):
        self.close()
        self.fix = C_kridi_fix()
        self.fix.show()

    def refresh_(self):
        while self.fix_table.rowCount() > 0 :
            self.articles_fix_tablewill_end.removeRow(0)

        curs.execute('SELECT  date, name , payed_money, rest FROM fix_C_kridi_history ORDER BY ID DESC')
        rus_ = curs.fetchall()
        self.fix_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.fix_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.fix_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

class C_kridi_fix(QWidget, fix_kridi_win_dir):# DONE
    def __init__(self, parent = None):
        super(C_kridi_fix, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('استخلاص قرض')
        self.fill_comboB()
        self.fix_kridi_done.clicked.connect(self.done)
        self.fix_kridi_save.clicked.connect(self.save_)
        self.fix_kridi_editeLine.textChanged.connect(self.check_state)
        self.fix_kridi_save.setEnabled(False)
        self.fix_kridi_editeLine.setEnabled(False)
        self.fix_kridi_editeLine.setValidator(QtGui.QIntValidator(1, 2147483647, self))
        self.fix_kridi_combo.currentTextChanged.connect(self.on_combo_changed)
        self.fix_kridi_history_btn.clicked.connect(self.fix_C_kridi_history)

    def fix_C_kridi_history(self):
        self.close()
        self.fixC_H = C_kridi_fix_history()
        self.fixC_H.show()

    def on_combo_changed(self):
        if self.fix_kridi_combo.currentText() != '':
            self.fix_kridi_editeLine.setEnabled(True)
        else:
            self.fix_kridi_editeLine.setEnabled(False)

        self.fix_kridi_editeLine.setText('')
        self.fix_kridi_editeLine.setFocus()

    def check_state(self):
        color = '#b6b6b6'#normal
        tt = 0
        for i in self.fix_kridi_editeLine.text():
            tt += int(i)
        curs.execute('SELECT total_rest FROM C_kridi WHERE name LIKE "{}"'.format(self.fix_kridi_combo.currentText()))
        rest = curs.fetchone()[0]
        if  self.fix_kridi_editeLine.text() == '':
            color = '#b6b6b6'#normal
            self.fix_kridi_save.setEnabled(False)
            self.rest_label.setText('')

            
        elif float(self.fix_kridi_editeLine.text()) > rest or tt == 0:
            color = '#f6989d'#red
            self.fix_kridi_save.setEnabled(False)
            self.rest_label.setText('خطأ')

        else:
            self.fix_kridi_save.setEnabled(True)
            self.rest_label.setText(str(rest - float(self.fix_kridi_editeLine.text())) + ' DH')
        self.fix_kridi_editeLine.setStyleSheet('QLineEdit { background-color: %s }' % color)

    def fill_comboB(self):
        self.fix_kridi_combo.clear()
        curs.execute('SELECT name FROM C_kridi')
        names = ['']
        for i in curs.fetchall():
            names.append(i[0])
        self.fix_kridi_combo.addItems(names)

    def done(self):
        self.close()
        self.home = Home()
        self.home.show()

    def save_(self):
        try:
            cll = self.fix_kridi_combo.currentText()
            curs.execute('UPDATE C_kridi SET total_recived = total_recived + {} WHERE name LIKE "{}"'
                    .format(self.fix_kridi_editeLine.text(), self.fix_kridi_combo.currentText()))
            curs.execute('UPDATE clients SET total_recived = total_recived + {} WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
                    .format(self.fix_kridi_editeLine.text(), cll.split(' ')[0], cll.split(' ')[1]))
            curs.execute('UPDATE C_kridi SET total_rest = debte - total_recived ')
            curs.execute('UPDATE clients SET total_rest = total_debted - total_recived ')
            curs.execute('SELECT total_rest FROM C_kridi WHERE name LIKE "{}"'.format(self.fix_kridi_combo.currentText()))
            curs.execute('INSERT INTO fix_C_kridi_history (date, name, payed_money, rest) VALUES ("{}", "{}", {}, {})'
                    .format(str(today), self.fix_kridi_combo.currentText(), self.fix_kridi_editeLine.text(), curs.fetchone()[0]))
            conn.commit()
        except ERROR as er:
            print(er)
            
class S_kridi_fix_history(QWidget, fix_kridi_history_win_dir):
    def __init__(self, parent = None):
        super(S_kridi_fix_history, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('القروض المستخلصة')
        self.refresh_()
        self.back_btn.clicked.connect(self.back__)
        self.print_btn.clicked.connect(self.pprint)




    def pprint(self):#TODO PRINT data 
        pass

    def back__(self):
        self.close()
        self.fix = S_kridi_fix()
        self.fix.show()

    def refresh_(self):
        while self.fix_table.rowCount() > 0 :
            self.articles_fix_tablewill_end.removeRow(0)

        curs.execute('SELECT  date, name , payed_money, rest FROM fix_S_kridi_history ORDER BY ID DESC')
        rus_ = curs.fetchall()
        self.fix_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.fix_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.fix_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

class S_kridi_fix(QWidget, fix_kridi_win_dir):
    def __init__(self, parent = None):
        super(S_kridi_fix, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('دفع قرض')
        self.fill_comboB()
        self.fix_kridi_done.clicked.connect(self.done)
        self.fix_kridi_save.clicked.connect(self.save_)
        self.fix_kridi_editeLine.textChanged.connect(self.check_state)
        self.fix_kridi_save.setEnabled(False)
        self.fix_kridi_editeLine.setEnabled(False)
        self.fix_kridi_editeLine.setValidator(QtGui.QIntValidator(1, 2147483647, self))
        self.fix_kridi_combo.currentTextChanged.connect(self.on_combo_changed)
        self.fix_kridi_history_btn.clicked.connect(self.fix_S_kridi_history)

    def fix_S_kridi_history(self):
        self.close()
        self.fixS_H = S_kridi_fix_history()
        self.fixS_H.show()

    def on_combo_changed(self):
        if self.fix_kridi_combo.currentText() != '':
            self.fix_kridi_editeLine.setEnabled(True)
        else:
            self.fix_kridi_editeLine.setEnabled(False)

        self.fix_kridi_editeLine.setText('')
        self.fix_kridi_editeLine.setFocus()
        

    def check_state(self):
        color = '#b6b6b6'#normal
        tt = 0
        for i in self.fix_kridi_editeLine.text():
            tt += int(i)
        curs.execute('SELECT total_rest FROM S_kridi WHERE name LIKE "{}"'.format(self.fix_kridi_combo.currentText()))
        rest = curs.fetchone()[0]
        if  self.fix_kridi_editeLine.text() == '':
            color = '#b6b6b6'#normal
            self.fix_kridi_save.setEnabled(False)
            self.rest_label.setText('')

            
        elif float(self.fix_kridi_editeLine.text()) > rest or tt == 0:
            color = '#f6989d'#red
            self.fix_kridi_save.setEnabled(False)
            self.rest_label.setText('خطأ')

        else:
            self.fix_kridi_save.setEnabled(True)
            self.rest_label.setText(str(rest - float(self.fix_kridi_editeLine.text())) + ' DH')
        self.fix_kridi_editeLine.setStyleSheet('QLineEdit { background-color: %s }' % color)

    def fill_comboB(self):
        self.fix_kridi_combo.clear()
        curs.execute('SELECT name FROM S_kridi')
        names = ['']
        for i in curs.fetchall():
            names.append(i[0])
        self.fix_kridi_combo.addItems(names)

    def done(self):
        self.close()
        self.home = Home()
        self.home.show()

    def save_(self):
        try:
            cll = self.fix_kridi_combo.currentText()
            curs.execute('UPDATE S_kridi SET total_recived = total_recived + {} WHERE name LIKE "{}"'
                    .format(self.fix_kridi_editeLine.text(), self.fix_kridi_combo.currentText()))
            curs.execute('UPDATE sellers SET total_recived = total_recived + {} WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
                    .format(self.fix_kridi_editeLine.text(), cll.split(' ')[0], cll.split(' ')[1]))
            curs.execute('UPDATE S_kridi SET total_rest = debte - total_recived ')
            curs.execute('UPDATE sellers SET total_rest = total_debted - total_recived ')
            curs.execute('SELECT total_rest FROM S_kridi WHERE name LIKE "{}"'.format(self.fix_kridi_combo.currentText()))
            curs.execute('INSERT INTO fix_S_kridi_history (date, name, payed_money, rest) VALUES ("{}", "{}", {}, {})'
                    .format(str(today), self.fix_kridi_combo.currentText(), self.fix_kridi_editeLine.text(), curs.fetchone()[0]))
            conn.commit()
        except ERROR as er:
            print(er)      
 
class C_kridi_history(QWidget, kridi_history_win_dir):
    def __init__(self, parent = None):
        super(C_kridi_history, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('القروض السابقة')
        self.back_from_history.clicked.connect(self.bback)
        self.print_history.clicked.connect(self.pprint)
        self.refresh()

    def refresh(self):
        while self.kridi_history_table.rowCount() > 0 :
            self.kridi_history_table.removeRow(0)

        curs.execute('SELECT date, name, article , qt, total, pay_date FROM C_kridi_history ORDER BY ID DESC')
        rus_ = curs.fetchall()
        self.kridi_history_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.kridi_history_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.kridi_history_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

    def pprint(self):#TODO PRINT data 
        pass

    def bback(self):
        self.close()
        self.home = Home()
        self.home.show()

class S_kridi_history(QWidget, kridi_history_win_dir):
    def __init__(self, parent = None):
        super(S_kridi_history, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('القروض السابقة')
        self.back_from_history.clicked.connect(self.bback)
        self.print_history.clicked.connect(self.pprint)
        self.refresh()

    def refresh(self):
        while self.kridi_history_table.rowCount() > 0 :
            self.kridi_history_table.removeRow(0)

        curs.execute('SELECT date, name, article , qt, total, pay_date FROM S_kridi_history ORDER BY ID DESC')
        rus_ = curs.fetchall()
        self.kridi_history_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.kridi_history_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.kridi_history_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))




    def pprint(self):#TODO PRINT data 
        pass


    def bback(self):
        self.close()
        self.home = Home()
        self.home.show()

class Edite_client(QWidget, edite_pea_win_dir):
    def __init__(self, parent = None):
        super(Edite_client, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('تعديل')
        self.done.clicked.connect(self.DONE)
        self.cliients_names_combo.currentTextChanged.connect(self.selected)
        self.refresh()
        self.save_.clicked.connect(self.save)


    def save(self):
        curs.execute('UPDATE clients SET F_name = "{}", L_name = "{}", cne = "{}", note = "{}" WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
        .format(self.F_name.text(), self.L_name.text(), self.cne.text(), self.note.toPlainText(), self.cliients_names_combo.currentText().split('|')[0], 
        self.cliients_names_combo.currentText().split('|')[1]))
        conn.commit()
        self.cliients_names_combo.setCurrentIndex(0)
        self.refresh()

    def selected(self):
        if self.cliients_names_combo.currentText() == '':
            self.F_name.setEnabled(False)
            self.L_name.setEnabled(False)
            self.cne.setEnabled(False)
            self.note.setEnabled(False)
            self.save_.setEnabled(False)

            self.F_name.setText('')
            self.L_name.setText('')
            self.cne.setText('')
            self.note.setText('')
        else:
            curs.execute('SELECT F_name, L_name, cne, note FROM clients WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
                .format(self.cliients_names_combo.currentText().split('|')[0], self.cliients_names_combo.currentText().split('|')[1]))
            info = curs.fetchone()
            print(info)

            self.F_name.setEnabled(True)
            self.L_name.setEnabled(True)
            self.cne.setEnabled(True)
            self.note.setEnabled(True)
            self.save_.setEnabled(True)

            self.F_name.setText(info[0])
            self.L_name.setText(info[1])
            self.cne.setText(info[2])
            self.note.setText(info[3])


    def refresh(self):
        self.cliients_names_combo.clear()
        curs.execute('SELECT F_name, L_name FROM clients')
        names =['']
        for i in curs.fetchall():
            names.append(i[0] + '|' + i[1])
        
        self.cliients_names_combo.addItems(names)



    def DONE(self):
        self.close()
        self.home = Home()
        self.home.show()

class Edite_seller(QWidget, edite_pea_win_dir):
    def __init__(self, parent = None):
        super(Edite_seller, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('تعديل')
        self.done.clicked.connect(self.DONE)
        self.cliients_names_combo.currentTextChanged.connect(self.selected)
        self.refresh()
        self.save_.clicked.connect(self.save)


    def save(self):
        curs.execute('UPDATE sellers SET F_name = "{}", L_name = "{}", cne = "{}", note = "{}" WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
        .format(self.F_name.text(), self.L_name.text(), self.cne.text(), self.note.toPlainText(), self.cliients_names_combo.currentText().split('|')[0], 
        self.cliients_names_combo.currentText().split('|')[1]))
        conn.commit()
        self.cliients_names_combo.setCurrentIndex(0)
        self.refresh()

    def selected(self):
        if self.cliients_names_combo.currentText() == '':
            self.F_name.setEnabled(False)
            self.L_name.setEnabled(False)
            self.cne.setEnabled(False)
            self.note.setEnabled(False)
            self.save_.setEnabled(False)

            self.F_name.setText('')
            self.L_name.setText('')
            self.cne.setText('')
            self.note.setText('')
        else:
            curs.execute('SELECT F_name, L_name, cne, note FROM sellers WHERE F_name LIKE "{}" AND L_name LIKE "{}"'
                .format(self.cliients_names_combo.currentText().split('|')[0], self.cliients_names_combo.currentText().split('|')[1]))
            info = curs.fetchone()
            print(info)

            self.F_name.setEnabled(True)
            self.L_name.setEnabled(True)
            self.cne.setEnabled(True)
            self.note.setEnabled(True)
            self.save_.setEnabled(True)

            self.F_name.setText(info[0])
            self.L_name.setText(info[1])
            self.cne.setText(info[2])
            self.note.setText(info[3])


    def refresh(self):
        self.cliients_names_combo.clear()
        curs.execute('SELECT F_name, L_name FROM sellers')
        names =['']
        for i in curs.fetchall():
            names.append(i[0] + '|' + i[1])
        
        self.cliients_names_combo.addItems(names)



    def DONE(self):
        self.close()
        self.home = Home()
        self.home.show()

class Edite_art(QWidget, edite_art_win_dir):
    def __init__(self, parent = None):
        super(Edite_art, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('تعديل')
        self.done.clicked.connect(self.DONE)
        self.articles_combo.currentTextChanged.connect(self.selected)
        self.refresh()
        self.save_.clicked.connect(self.save)


    def save(self):
        curs.execute('UPDATE articles SET name = "{}", type = "{}", qt = {}, price = {}, note = "{}" WHERE name LIKE "{}" '
        .format(self.name.text(), self.type.text(), self.qt.text(), self.price.text(), self.note.toPlainText(), self.articles_combo.currentText() ))
        conn.commit()
        self.articles_combo.setCurrentIndex(0)
        self.refresh()

    def selected(self):
        if self.articles_combo.currentText() == '':
            self.name.setEnabled(False)
            self.type.setEnabled(False)
            self.qt.setEnabled(False)
            self.price.setEnabled(False)
            self.note.setEnabled(False)
            self.save_.setEnabled(False)

            self.name.setText('')
            self.type.setText('')
            self.qt.setText('')
            self.price.setText('')
            self.note.setText('')
        else:
            curs.execute('SELECT name, type, qt, price, note FROM articles WHERE name LIKE "{}" '
                .format(self.articles_combo.currentText()))
            info = curs.fetchone()
            print(info)

            self.name.setEnabled(True)
            self.type.setEnabled(True)
            self.qt.setEnabled(True)
            self.price.setEnabled(True)
            self.note.setEnabled(True)
            self.save_.setEnabled(True)

            self.name.setText(info[0])
            self.type.setText(info[1])
            self.qt.setText(str(info[2]))
            self.price.setText(str(info[3]))
            self.note.setText(info[4])


    def refresh(self):
        self.articles_combo.clear()
        curs.execute('SELECT name FROM articles')
        names =['']
        for i in curs.fetchall():
            names.append(i[0])
        
        self.articles_combo.addItems(names)



    def DONE(self):
        self.close()
        self.home = Home()
        self.home.show()

class Home(QWidget, home_win_dir):# almost...
    def __init__(self, parent = None):
        super(Home, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.logOut_Button.clicked.connect(self.log_out)
        self.refresh_data()
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.home_tabWidget.currentChanged.connect(self.onTabChange)
        self.setWindowTitle('الرئيسية')
        self.sell_choose_article_comboBox.currentTextChanged.connect(self.set_sell_article_type)
        self.save_selling_operation_pushButton.clicked.connect(self.save_sell_operation)
        self.save_buying_operation_pushButton.clicked.connect(self.save_buy_operation)
        self.sell_qt.lineEdit().setEnabled(False)
        self.sell_qt.valueChanged.connect(self.on_sell_qt_spin_box_changed)
        self.buy_article_qt.valueChanged.connect(self.on_buy_qt_spin_box_changed)
        self.sell_date_to_pay.setDateTime(QtCore.QDateTime.currentDateTime())
        self.add_client_Button.clicked.connect(self.add_new_client)
        self.add_seller_Button.clicked.connect(self.add_new_seller)
        self.add_art_Button.clicked.connect(self.add_new_article)
        self.buy_articl_name_lineEdit.hide()
        self.buy_choose_article_comboBox.hide()
        self.buy_new_article.clicked.connect(self.new_art_)
        self.buy_exists_article.clicked.connect(self.exists_art_)
        self.reset_buying_info_pushButton.clicked.connect(self.rest_btns)
        self.buy_art_stat = ''
        self.buy_choose_article_comboBox.currentTextChanged.connect(self.set_bb_article_type)
        self.sell_history_Button.clicked.connect(self.open_sell_history)
        self.buy_and_buy_history_Button.clicked.connect(self.open_buy_history)
        self.exit_Button.clicked.connect(self.exit)
        self.remove_client_Button.clicked.connect(self.delete_client)
        self.remove_seller_Button.clicked.connect(self.remove_seller)
        self.take_kridi_btn.clicked.connect(self.fix_C_kridi)
        self.pay_kridi_btn.clicked.connect(self.fix_S_kridi)
        self.C_kridi_history_btn.clicked.connect(self.c_k_history)
        self.S_kridi_history_btn.clicked.connect(self.s_k_history)
        self.edit_client_Button.clicked.connect(self.edit_client)
        self.edit_seller_Button.clicked.connect(self.edit_seller)
        self.edit_art_Button.clicked.connect(self.edit_article)
        self.print_clients_table.clicked.connect(self.print_c_info)###
        self.searsh_client_EditText.textChanged.connect(self.searsh_clients)
        self.searsh_seller_EditText.textChanged.connect(self.searsh_sellers)
        self.searsh_art.textChanged.connect(self.searsh_articles)
        self.setting_Button.clicked.connect(self.settt)
        self.print_sellers_info.clicked.connect(self.print_s_info)


    def searsh_articles(self):
        # self.searsh_art.setText('*' + self.searsh_art.text() + '*')
        if self.searsh_art.text() != '' or self.searsh_art.text() != ' ':
            #sellers table
            while self.articles_table.rowCount() > 0 :
                self.articles_table.removeRow(0)
            curs.execute('''
            SELECT name,  type, qt, price, note FROM articles 
                WHERE name LIKE "{}" OR type LIKE "{}" OR qt = "{}" OR price = "{}" OR note LIKE "{}"'''
            .format( '%' + str(self.searsh_art.text()) + '%', '%' + str(self.searsh_art.text()) + '%', '%' + self.searsh_art.text() + '%', '%' + self.searsh_art.text() + '%',
            '%' + self.searsh_art.text() + '%'))
            # print(curs.fetchall())
            rus__ = curs.fetchall()
            self.articles_table.setRowCount(0)
            for r_n, r_d in enumerate(rus__):
                self.articles_table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.articles_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))


        else:

            #sellers table
            while self.articles_table.rowCount() > 0 :
                self.articles_table.removeRow(0)
            curs.execute(
                'SELECT  name,  type, qt, price, note  FROM articles ORDER BY ID DESC;')
            rus__ = curs.fetchall()
            self.articles_table.setRowCount(0)
            for r_n, r_d in enumerate(rus__):
                self.articles_table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.articles_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

    def searsh_sellers(self):
        # self.searsh_seller_EditText.setText('*' + self.searsh_seller_EditText.text() + '*')
        if self.searsh_seller_EditText.text() != '' or self.searsh_seller_EditText.text() != ' ':
            #sellers table
            while self.sellers_table.rowCount() > 0 :
                self.sellers_table.removeRow(0)
            curs.execute('''
            SELECT F_name,  L_name, cne, total_debted, total_recived, total_rest, pay_date, note FROM sellers 
                WHERE F_name LIKE "{}" OR L_name LIKE "{}" OR cne LIKE "{}" OR total_debted = "{}" OR total_recived = "{}" OR pay_date LIKE "{}" OR note LIKE "{}"'''
            .format('%' + self.searsh_seller_EditText.text() + '%', '%' + self.searsh_seller_EditText.text() + '%', '%' + self.searsh_seller_EditText.text() + '%',
            '%' + self.searsh_seller_EditText.text() + '%', '%' + str(self.searsh_seller_EditText.text()) + '%', '%' + str(self.searsh_seller_EditText.text()) + '%', 
            '%' + self.searsh_seller_EditText.text() + '%'))
            # print(curs.fetchall())
            rus__ = curs.fetchall()
            self.sellers_table.setRowCount(0)
            for r_n, r_d in enumerate(rus__):
                self.sellers_table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.sellers_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))


        else:

            #sellers table
            while self.sellers_table.rowCount() > 0 :
                self.sellers_table.removeRow(0)
            curs.execute(
                'SELECT  F_name,  L_name, cne, total_debted, total_recived, total_rest, pay_date, note FROM sellers ORDER BY ID DESC;')
            rus__ = curs.fetchall()
            self.sellers_table.setRowCount(0)
            for r_n, r_d in enumerate(rus__):
                self.sellers_table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.sellers_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

    def searsh_clients(self):
        # self.searsh_client_EditText.setText('*' + self.searsh_client_EditText.text() + '*')
        if self.searsh_client_EditText.text() != '' or self.searsh_client_EditText.text() != ' ':
            #clients table
            while self.clients_table.rowCount() > 0 :
                self.clients_table.removeRow(0)
            curs.execute('''
            SELECT F_name,  L_name, cne, total_debted, total_recived, total_rest, pay_date, note FROM clients 
                WHERE F_name LIKE "{}" OR L_name LIKE "{}" OR cne LIKE "{}" OR total_debted = "{}" OR total_recived = "{}" OR pay_date LIKE "{}" OR note LIKE "{}"'''
            .format('%' + self.searsh_client_EditText.text() + '%', '%' + self.searsh_client_EditText.text() + '%', '%' + self.searsh_client_EditText.text() + '%',
            '%' + self.searsh_client_EditText.text() + '%', '%' + str(self.searsh_client_EditText.text()) + '%', '%' + str(self.searsh_client_EditText.text()) + '%', 
            '%' + self.searsh_client_EditText.text() + '%'))
            # print(curs.fetchall())
            rus__ = curs.fetchall()
            self.clients_table.setRowCount(0)
            for r_n, r_d in enumerate(rus__):
                self.clients_table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.clients_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))


        else:

            #clients table
            while self.clients_table.rowCount() > 0 :
                self.clients_table.removeRow(0)
            curs.execute(
                'SELECT  F_name,  L_name, cne, total_debted, total_recived, total_rest, pay_date, note FROM clients ORDER BY ID DESC;')
            rus__ = curs.fetchall()
            self.clients_table.setRowCount(0)
            for r_n, r_d in enumerate(rus__):
                self.clients_table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.clients_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))


    #TODO PRINT 
    def print_s_info(self):
        pass


    #TODO PRINT 
    def print_c_info(self):
        pass

   
    def c_k_history(self):
        self.c_k_h = C_kridi_history()
        self.c_k_h.show()
        self.close()

    def s_k_history(self):
        self.c_k_h = S_kridi_history()
        self.c_k_h.show()
        self.close()

    def fix_C_kridi(self):
        self.close()
        self.fix_C_K = C_kridi_fix()
        self.fix_C_K.show()

    def fix_S_kridi(self):
        self.close()
        self.fix_S_K = S_kridi_fix()
        self.fix_S_K.show()

    def delete_client(self):
        self.close()
        self.delete_client_win = Remove_client()
        self.delete_client_win.show()

    def exit(self):
        conn.close()
        self.close()

    def set_bb_article_type(self):

        try:
            curs.execute(
                'SELECT type FROM articles WHERE name LIKE "{}";'.format(
                    self.buy_choose_article_comboBox.currentText()))
            self.buy_article_type_lineEdit.setText(curs.fetchone()[0])
        except Exception:
            print(Exception)

        curs.execute('SELECT price FROM articles WHERE name LIKE "{}"'.format(self.buy_choose_article_comboBox.currentText()))
        self.buy_article_price.setValue(curs.fetchone()[0])

    def new_art_(self):
        self.buy_art_stat = 'new'
        self.buy_articl_name_lineEdit.show()
        self.buy_new_article.hide()
        self.buy_exists_article.hide()
        self.buy_article_price.setEnabled(True)
        self.buy_article_type_lineEdit.setEnabled(True)

    def exists_art_(self):
        self.buy_art_stat = 'exists'
        self.buy_choose_article_comboBox.show()
        self.buy_exists_article.hide()
        self.buy_new_article.hide()
        self.buy_article_price.setEnabled(False)
        self.buy_article_type_lineEdit.setEnabled(False)
        self.buy_choose_article_comboBox.clear()
        self.articles_list_ = ['']
        curs.execute('SELECT name FROM articles WHERE qt > 0')
        for i in curs.fetchall():
            self.articles_list_.append(i[0])
        print(self.articles_list_)
        self.buy_choose_article_comboBox.addItems(self.articles_list_)

    def rest_btns(self):
        self.buy_art_stat = ''
        self.buy_articl_name_lineEdit.hide()
        self.buy_choose_article_comboBox.hide()
        self.buy_exists_article.show()
        self.buy_new_article.show()
        self.buy_article_price.setEnabled(True)
        self.buy_article_type_lineEdit.setEnabled(True)

    def log_out(self):
        curs.execute('SELECT first FROM tools')
        if curs.fetchone()[0] == 1:
            self.close()
            self.log_out = VirificationAlert()
            self.log_out.show()
        else:
            self.exit()

    def refresh_data(self):  # refresh all app data
        self.refresh_labels()
        self.refresh_tables()
        self.fill_combos()


    def onTabChange(self): #TODO Fix this shit
        print('the current tab is : ', self.home_tabWidget.currentIndex())
        self.refresh_data()
        if self.home_tabWidget.currentIndex() == 0:
            # current_tab = 0
            self.setWindowTitle('الرئيسية')

        elif self.home_tabWidget.currentIndex() == 1:
            # current_tab = 1
            self.setWindowTitle('المعاملات')

        elif self.home_tabWidget.currentIndex() == 2:
            # current_tab = 2
            self.setWindowTitle('القروض')

        elif self.home_tabWidget.currentIndex() == 3:
            # current_tab = 3
            self.setWindowTitle('الزبائن')

        elif self.home_tabWidget.currentIndex() == 4:
            # current_tab = 4
            self.setWindowTitle('الموزعين')

        elif self.home_tabWidget.currentIndex() == 5:
            # current_tab = 4
            self.setWindowTitle('السلع')

    def refresh_labels(self):
        curs.execute('SELECT NAME FROM user')
        self.wilcom_label.setText('مرحبا بالسيد ' + curs.fetchone()[0])
        self.date_time_label.setText(str(today))
        curs.execute('SELECT hadit FROM hadit ;')
        self.hadit_list = []
        for i in curs.fetchall():
            self.hadit_list.append(i[0])
        self.hadit_label.setText(random.choice(self.hadit_list))
        # print(self.hadit_list)
        curs.execute('SELECT sum(total_rest) FROM C_kridi')
        total_rest_C = curs.fetchone()[0]
        
        self.total_money_debt.setText(str(total_rest_C))
        self.total_money_debt.setText(self.total_money_debt.text() + ' DH ')
        curs.execute('SELECT sum(total_rest) FROM S_kridi ')
        total_debt_S = curs.fetchone()[0]
        
        self.money_u_have_to_pay.setText(str(total_debt_S))
        self.money_u_have_to_pay.setText(self.money_u_have_to_pay.text() + ' DH ')
        print( 'set_total_sellers_dept_money', total_debt_S)

        curs.execute('SELECT COUNT(ID) FROM clients WHERE total_rest > 0;')
        self.clients_debteds_counter.setText(str(curs.fetchone()[0]))
        curs.execute('SELECT COUNT(ID) FROM clients WHERE pay_date LIKE "{}"'.format(str(today)))
        # print(curs.fetchone()[0] == str(today))
        clients_pay_today_list = curs.fetchone()[0]
        
        self.clients_must_pay_counter_today.setText(str(clients_pay_today_list))

    def refresh_tables(self):
        curs.execute('UPDATE clients SET total_rest = total_debted - total_recived')
        curs.execute('UPDATE sellers SET total_rest = total_debted - total_recived')
        curs.execute('DELETE FROM C_kridi WHERE total_rest = 0')
        curs.execute('DELETE FROM S_kridi WHERE total_rest = 0')
        curs.execute('DELETE FROM articles WHERE qt = 0')
        conn.commit()
        #clients have to pay today table
        while self.clients_pay_today_table.rowCount() > 0 :
            self.clients_pay_today_table.removeRow(0)

        curs.execute('SELECT  name , total_rest FROM C_kridi  WHERE pay_date LIKE "{}"'.format(str(today)))
        rus_ = curs.fetchall()
        self.clients_pay_today_table.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.clients_pay_today_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.clients_pay_today_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        print('rus_- = ', rus_)


        #articles will end table
        while self.articles_will_end.rowCount() > 0 :
            self.articles_will_end.removeRow(0)

        curs.execute('SELECT  name , qt FROM articles  WHERE qt < 10 ')
        rus_ = curs.fetchall()
        self.articles_will_end.setRowCount(0)
        for r_n, r_d in enumerate(rus_):
            self.articles_will_end.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.articles_will_end.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))



        #clients table
        while self.clients_table.rowCount() > 0 :
            self.clients_table.removeRow(0)
        curs.execute(
            'SELECT  F_name,  L_name, cne, total_debted, total_recived, total_rest, pay_date, note FROM clients ORDER BY ID DESC;')
        rus__ = curs.fetchall()
        self.clients_table.setRowCount(0)
        for r_n, r_d in enumerate(rus__):
            self.clients_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.clients_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        print(curs.fetchall())
        #
        #TODO : contuni with all tables
        #sellers table
        while self.sellers_table.rowCount() > 0 :
            self.sellers_table.removeRow(0)

        curs.execute(
            'SELECT  F_name,  L_name, cne, total_debted, total_recived, total_rest, pay_date, note FROM sellers ORDER BY ID DESC;')
        rus__ = curs.fetchall()
        self.sellers_table.setRowCount(0)
        for r_n, r_d in enumerate(rus__):
            self.sellers_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.sellers_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))


        #articles table

        while self.articles_table.rowCount() > 0 :
            self.articles_table.removeRow(0)
        curs.execute(
            'SELECT name, type, qt, price, note FROM articles ORDER BY ID DESC;')
        rus__ = curs.fetchall()
        self.articles_table.setRowCount(0)
        for r_n, r_d in enumerate(rus__):
            self.articles_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.articles_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        while self.C_kridi_table.rowCount() > 0 :
            self.C_kridi_table.removeRow(0)
        curs.execute(
            'SELECT name, cne, total_rest FROM C_kridi ORDER BY ID DESC;')
        rus__ = curs.fetchall()
        self.C_kridi_table.setRowCount(0)
        for r_n, r_d in enumerate(rus__):
            self.C_kridi_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.C_kridi_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))


        while self.S_kridi_table.rowCount() > 0 :
            self.S_kridi_table.removeRow(0)
        curs.execute(
            'SELECT name, cne, total_rest FROM S_kridi ORDER BY ID DESC;')
        rus__ = curs.fetchall()
        self.S_kridi_table.setRowCount(0)
        for r_n, r_d in enumerate(rus__):
            self.S_kridi_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.S_kridi_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))



        #clients kridi
        while self.S_kridi_table.rowCount() > 0 :
            self.S_kridi_table.removeRow(0)
        curs.execute(
            'SELECT name, cne, total_rest FROM S_kridi ORDER BY ID DESC;')
        rus__ = curs.fetchall()
        self.S_kridi_table.setRowCount(0)
        for r_n, r_d in enumerate(rus__):
            self.S_kridi_table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.S_kridi_table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

    def fill_combos(self):

        self.sell_choose_article_comboBox.clear()
        self.sell_choose_buyer_comboBox.clear()
        self.buy_choose_seller_comboBox.clear()
        self.articles_list_  = ['']
        curs.execute('SELECT name FROM articles WHERE qt > 0')
        for i in curs.fetchall():
            self.articles_list_.append(i[0])
        print(self.articles_list_)
        self.sell_choose_article_comboBox.addItems(self.articles_list_)

        clients_names_list_ = []

        curs.execute('SELECT F_name, L_name  FROM clients')
        for fn, ln in curs.fetchall():

            clients_names_list_.append(fn + ' ' + ln)
        print('CLIENTS : ', clients_names_list_)
        self.sell_choose_buyer_comboBox.addItems(clients_names_list_)

        selers_names_list_ = ['']
        curs.execute('SELECT F_name, L_name FROM sellers ')
        for sfn, sln in curs.fetchall():
            selers_names_list_.append(sfn + ' ' + sln)
        print('SELLRES : ', selers_names_list_)
        self.buy_choose_seller_comboBox.addItems(selers_names_list_)

    def set_sell_article_type(self):
        print(str(self.sell_choose_article_comboBox.currentText()))
        # r_count = curs.execute('SELECT COUNT(*) FROM articles')
        # print('articles table has {} rows '.format(str(r_count.fetchone())))
        if self.sell_choose_article_comboBox.currentText() == '' or self.sell_choose_article_comboBox.currentText() == ' ':
            self.sell_article_type.setText('')
        else:
            self.sell_qt.lineEdit().setEnabled(True)
            curs.execute('SELECT type FROM articles WHERE name like "{}"'.format(
                str(self.sell_choose_article_comboBox.currentText())))
            self.sell_article_type.setText(curs.fetchone()[0])
            curs.execute('SELECT qt FROM articles WHERE name like "{}"'.format(
                str(self.sell_choose_article_comboBox.currentText())))
            self.sell_qt.setMaximum(curs.fetchone()[0])

    def on_sell_qt_spin_box_changed(self):
        if str(self.sell_choose_article_comboBox.currentText()) != '' or str(self.sell_choose_article_comboBox.currentText()) != '':
            curs.execute('SELECT price FROM articles WHERE name like "{}"'.format(
                    str(self.sell_choose_article_comboBox.currentText())))
            total_price = self.sell_qt.value() * curs.fetchone()[0]
            self.total_selling_price.setText(str(total_price) + ' DH')
        else:
            self.total_selling_price.setText('إختر سلعة لإظهار المجموع')

    def on_buy_qt_spin_box_changed(self):
        if self.buy_art_stat == 'exists':
            if str(self.buy_choose_article_comboBox.currentText()) != '' or str(self.buy_choose_article_comboBox.currentText()) != '':
                curs.execute('SELECT price FROM articles WHERE name like "{}"'.format(
                        str(self.buy_choose_article_comboBox.currentText())))
                total_price = self.buy_article_qt.value() * curs.fetchone()[0]
                self.total_buying_price.setText(str(total_price) + ' DH')
            else:
                self.total_buying_price.setText('إختر سلعة لإظهار المجموع')
        elif self.buy_art_stat == 'new':
            total_price = self.buy_article_qt.value() * self.buy_article_price.value()
            self.total_buying_price.setText(str(total_price) + ' DH')

    def save_sell_operation(self):
        if str(self.sell_choose_article_comboBox.currentText()) == '' or  str(self.sell_choose_article_comboBox.currentText()) == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إختيار سلعة ')
            self.err.setWindowTitle('خطأ')
            self.sell_choose_article_comboBox.setFocus()
        elif str(self.sell_choose_buyer_comboBox.currentText()) == '' or  str(self.sell_choose_buyer_comboBox.currentText()) == ' ':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إختيار زبون')
            self.err.setWindowTitle('خطأ')
            self.sell_choose_buyer_comboBox.setFocus()
        elif self.sell_qt.value() == 0:
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('لايمكن بيع هذه الكمية')
            self.err.setWindowTitle('خطأ')
            self.sell_qt.setFocus()
        elif self.sell_pay_by_debt_radioButton.isChecked():
            curs.execute('SELECT price FROM articles WHERE name LIKE "{}" '.format(
                self.sell_choose_article_comboBox.currentText()))
            price = curs.fetchone()[0]

            curs.execute('SELECT name FROM C_kridi ')
            cll = []
            for n in curs.fetchall():
                cll.append(n[0])
            tt = self.sell_qt.value() * price
            if self.sell_choose_buyer_comboBox.currentText() in cll:
                curs.execute(
                    'UPDATE C_kridi SET debte = debte + {} , pay_date = "{}" WHERE name LIKE "{}"'
                        .format(tt, str(self.sell_date_to_pay.date().toPyDate()),
                                self.sell_choose_buyer_comboBox.currentText()))
                curs.execute('UPDATE C_kridi SET total_rest = debte - total_recived')
                conn.commit()
            else:
                cne = curs.execute('SELECT cne FROM clients WHERE F_name LIKE "{}" AND L_name LIKE "{}" '
                                   .format(self.sell_choose_buyer_comboBox.currentText().split(' ')[0],
                                           self.sell_choose_buyer_comboBox.currentText().split(' ')[1]))

                curs.execute('''
                                                 INSERT INTO C_kridi(name, cne, debte, total_recived, pay_date) VALUES("{}", "{}", {}, 0, "{}")
                                                              '''.format(self.sell_choose_buyer_comboBox.currentText(),
                                                                         cne.fetchone()[0], int(tt),
                                                                         self.sell_date_to_pay.date().toPyDate()))
                curs.execute('UPDATE C_kridi SET total_rest = debte - total_recived')

                conn.commit()
            curs.execute('''
            UPDATE articles 
            SET qt = qt - {}
             WHERE name LIKE "{}";
             '''.format( self.sell_qt.value()
                         , self.sell_choose_article_comboBox.currentText()))

            curs.execute('''
            SELECT price 
            FROM articles 
            WHERE name LIKE "{}";
             '''.format(self.sell_choose_article_comboBox.currentText()))

            curs.execute('''
            UPDATE clients SET
            total_debted = total_debted + {}
            WHERE F_name like "{}"
            AND L_name like "{}" ;
            '''.format(self.sell_qt.value() * curs.fetchone()[0]
                       , self.sell_choose_buyer_comboBox.currentText().split(' ')[0]
                       , self.sell_choose_buyer_comboBox.currentText().split(' ')[1] ))


            curs.execute('''
            UPDATE clients SET
            pay_date = "{}"
            WHERE F_name like "{}"
            AND L_name like "{}" ;
            '''.format(str(self.sell_date_to_pay.date().toPyDate())
                       , self.sell_choose_buyer_comboBox.currentText().split(' ')[0]
                       , self.sell_choose_buyer_comboBox.currentText().split(' ')[1] ))


            curs.execute('''
                                    SELECT price 
                                    FROM articles 
                                    WHERE name LIKE "{}";
                                     '''.format(self.sell_choose_article_comboBox.currentText()))
            curs.execute('''
                                    INSERT INTO selles_history (date, client, article, qt, not_payed_yet, pay_date) VALUES("{}", "{}", "{}", {}, {}, "{}") ;
                                    '''.format(today, self.sell_choose_buyer_comboBox.currentText(),
                                               self.sell_choose_article_comboBox.currentText(),
                                               self.sell_qt.value(), self.sell_qt.value() * curs.fetchone()[0],
                                               str(self.sell_date_to_pay.date().toPyDate())))


            conn.commit()
            print(self.sell_date_to_pay.date().toPyDate())
            self.sell_choose_article_comboBox.setCurrentIndex(0)
            self.sell_choose_buyer_comboBox.setCurrentIndex(0)
            self.sell_qt.setValue(0)
            self.refresh_data()
        else:
            curs.execute('''
                        UPDATE articles 
                        SET qt = qt - {}
                         WHERE name LIKE "{}";
                         '''.format(self.sell_qt.value()
                                    , self.sell_choose_article_comboBox.currentText()))

            curs.execute('''
                        SELECT price 
                        FROM articles 
                        WHERE name LIKE "{}";
                         '''.format(self.sell_choose_article_comboBox.currentText()))

            curs.execute('''
                        INSERT INTO selles_history (date, client, article, qt, payed, pay_date) VALUES("{}", "{}", "{}", {}, {}, "{}") 
                        '''.format(today, self.sell_choose_buyer_comboBox.currentText(), self.sell_choose_article_comboBox.currentText(),
                                   self.sell_qt.value(),self.sell_qt.value() * curs.fetchone()[0], str(self.sell_date_to_pay.date().toPyDate()) ))

            conn.commit()
            print('DONE')
            self.sell_choose_article_comboBox.setCurrentIndex(0)
            self.sell_choose_buyer_comboBox.setCurrentIndex(0)
            self.sell_qt.setValue(0)
            self.refresh_data()

    def save_buy_operation(self):

        print(self.buy_art_stat)
        if self.buy_art_stat == '':
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('المرجو إدخال إسم السلعة ')
            self.err.setWindowTitle('خطأ')

        if self.buy_art_stat == 'new':
            if self.buy_articl_name_lineEdit.text() == '' or self.buy_articl_name_lineEdit.text() == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('المرجو إدخال إسم السلعة ')
                self.err.setWindowTitle('خطأ')
                self.buy_articl_name_lineEdit.setFocus()

            elif self.buy_article_type_lineEdit.text() == '' or self.buy_article_type_lineEdit.text() == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('المرجو إدخال نوع السلعة ')
                self.err.setWindowTitle('خطأ')
                self.buy_article_type_lineEdit.setFocus()

            elif str(self.buy_choose_seller_comboBox.currentText()) == '' or str(
                    self.buy_choose_seller_comboBox.currentText()) == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('المرجو اختيار الموزع')
                self.err.setWindowTitle('خطأ')
                self.buy_choose_seller_comboBox.setFocus()

            elif self.buy_article_price.value() == 0:
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('لا يمكنك ترك الثمن 0 درهم')
                self.err.setWindowTitle('خطأ')
                self.buy_article_price.setFocus()

            elif self.buy_article_qt.value() == 0:
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('لا يمكنك ترك الكمية 0')
                self.err.setWindowTitle('خطأ')
                self.buy_article_price.setFocus()

            elif self.buy_pay_by_debt_radioButton.isChecked():

                ##########################
                # curs.execute('SELECT price FROM articles WHERE name LIKE "{}" '.format(self.sell_choose_article_comboBox.currentText()))
                price = self.buy_article_price.value()

                curs.execute('SELECT name FROM S_kridi ')
                Sll = []
                for n in curs.fetchall():
                    Sll.append(n[0])
                tt = self.buy_article_qt.value() * price
                if self.buy_choose_seller_comboBox.currentText() in Sll:
                    curs.execute(
                        'UPDATE S_kridi SET debte = debte + {} , pay_date = "{}" WHERE name LIKE "{}"'
                            .format(tt, str(self.buy_date_to_pay.date().toPyDate()),
                                    self.buy_choose_seller_comboBox.currentText()))
                    curs.execute('UPDATE S_kridi SET total_rest = debte - total_recived')
                    conn.commit()
                else:
                    cne = curs.execute('SELECT cne FROM sellers WHERE F_name LIKE "{}" AND L_name LIKE "{}" '
                                       .format(self.buy_choose_seller_comboBox.currentText().split(' ')[0],
                                               self.buy_choose_seller_comboBox.currentText().split(' ')[1]))

                    curs.execute('''
                                 INSERT INTO S_kridi(name, cne, debte, total_recived, pay_date) VALUES("{}", "{}", {}, 0, "{}")
                                 '''.format(
                        self.buy_choose_seller_comboBox.currentText(),
                        cne.fetchone()[0], int(tt),
                        self.buy_date_to_pay.date().toPyDate()))
                    curs.execute('UPDATE S_kridi SET total_rest = debte - total_recived')

                    conn.commit()
                ###############################

                curs.execute('''
                              INSERT INTO articles
                                          (name, type, qt, price) VALUES ("{}", "{}", {}, {});
                                          '''.format(self.buy_articl_name_lineEdit.text(),
                                                     self.buy_article_type_lineEdit.text(), self.buy_article_qt.value(),
                                                     self.buy_article_price.value()))

                curs.execute('UPDATE sellers SET total_debted = total_debted + {}, pay_date = "{}"  WHERE F_name LIKE "{}" AND L_name LIKE "{}";'.format(
                    self.buy_article_price.value() * self.buy_article_qt.value(), str(self.buy_date_to_pay.date().toPyDate())
                        , self.buy_choose_seller_comboBox.currentText().split(' ')[0], self.buy_choose_seller_comboBox.currentText().split(' ')[1]))

                curs.execute('INSERT INTO buys_history (date, seller, article, qt, not_payed_yet, pay_date) VALUES("{}", "{}", "{}", {}, {}, "{}") ;'
                             .format(today, self.buy_choose_seller_comboBox.currentText(), self.buy_articl_name_lineEdit.text(), self.buy_article_qt.value()
                                     , self.buy_article_price.value() * self.buy_article_qt.value(), str(self.buy_date_to_pay.date().toPyDate())))




            else:

                curs.execute('''INSERT INTO articles
                                 (name, type, qt, price) VALUES ("{}", "{}", {}, {});
                                                          '''.format(self.buy_articl_name_lineEdit.text(),
                                                                     self.buy_article_type_lineEdit.text(),
                                                                     self.buy_article_qt.value(),
                                                                     self.buy_article_price.value()))

                curs.execute(
                    'UPDATE sellers SET total_recived = total_recived + {}, pay_date = "{}"  WHERE F_name LIKE "{}" AND L_name LIKE "{}";'.format(
                        self.buy_article_price.value() * self.buy_article_qt.value(),
                        today
                        , self.buy_choose_seller_comboBox.currentText().split(' ')[0],
                        self.buy_choose_seller_comboBox.currentText().split(' ')[1]))

                curs.execute(
                    '''INSERT INTO buys_history
                     (date, seller, article, qt, payed, pay_date) 
                     VALUES("{}", "{}", "{}", {}, {}, "{}") ;
                     '''.format(today, self.buy_choose_seller_comboBox.currentText(), self.buy_articl_name_lineEdit.text(),
                            self.buy_article_qt.value()
                            , self.buy_article_price.value() * self.buy_article_qt.value(),
                            today))


            conn.commit()
        elif self.buy_art_stat == 'exists':
            if self.buy_choose_article_comboBox.currentText() == '' or self.buy_choose_article_comboBox.currentText() == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('المرجو إختيار إسم السلعة ')
                self.err.setWindowTitle('خطأ')
                self.buy_choose_article_comboBox.setFocus()


            elif self.buy_article_type_lineEdit.text() == '' or self.buy_article_type_lineEdit.text() == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('المرجو إدخال نوع السلعة ')
                self.err.setWindowTitle('خطأ')
                self.buy_article_type_lineEdit.setFocus()

            elif str(self.buy_choose_seller_comboBox.currentText()) == '' or str(self.buy_choose_seller_comboBox.currentText()) == ' ':
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('المرجو اختيار الموزع')
                self.err.setWindowTitle('خطأ')
                self.buy_choose_seller_comboBox.setFocus()

            # elif self.buy_article_price.value() == 0:
            #     self.err = QtWidgets.QErrorMessage()
            #     self.err.showMessage('لا يمكنك ترك الثمن 0 درهم')
            #     self.err.setWindowTitle('خطأ')
            #     self.buy_article_price.setFocus()

            elif self.buy_article_qt.value() == 0:
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('لا يمكنك ترك الكمية 0')
                self.err.setWindowTitle('خطأ')
                self.buy_article_price.setFocus()

            elif self.buy_pay_by_debt_radioButton.isChecked():

                ##########################
                curs.execute('SELECT price FROM articles WHERE name LIKE "{}" '.format(self.buy_choose_article_comboBox.currentText()))
                price = curs.fetchone()

                curs.execute('SELECT name FROM S_kridi ')
                Sll = []
                for n in curs.fetchall():
                    Sll.append(n[0])
                tt = self.buy_article_qt.value() * price
                if self.buy_choose_seller_comboBox.currentText() in Sll:
                    curs.execute(
                        'UPDATE S_kridi SET debte = debte + {} , pay_date = "{}" WHERE name LIKE "{}"'
                            .format(tt, str(self.buy_date_to_pay.date().toPyDate()),
                                    self.buy_choose_seller_comboBox.currentText()))
                    curs.execute('UPDATE S_kridi SET total_rest = debte - total_recived')
                    conn.commit()
                else:
                    cne = curs.execute('SELECT cne FROM sellers WHERE F_name LIKE "{}" AND L_name LIKE "{}" '
                                       .format(self.buy_choose_seller_comboBox.currentText().split(' ')[0],
                                               self.buy_choose_seller_comboBox.currentText().split(' ')[1]))

                    curs.execute('''
                                 INSERT INTO S_kridi(name, cne, debte, total_recived, pay_date) VALUES("{}", "{}", {}, 0, "{}")
                                 '''.format(
                        self.buy_choose_seller_comboBox.currentText(),
                        cne.fetchone()[0], int(tt),
                        self.buy_date_to_pay.date().toPyDate()))
                    curs.execute('UPDATE S_kridi SET total_rest = debte - total_recived')

                    conn.commit()
                ###############################


                curs.execute('''
                UPDATE articles SET qt = qt - {} WHERE name LIKE "{}"
                '''.format(self.buy_article_qt.value(), self.buy_choose_article_comboBox.currentText()))



                curs.execute(
                    'UPDATE sellers SET total_debted = total_debted + {}, pay_date = "{}"  WHERE F_name LIKE "{}" AND L_name LIKE "{}";'.format(
                        self.buy_article_price.value() * self.buy_article_qt.value(),
                        str(self.buy_date_to_pay.date().toPyDate())
                        , self.buy_choose_seller_comboBox.currentText().split(' ')[0],
                        self.buy_choose_seller_comboBox.currentText().split(' ')[1]))

                curs.execute(
                    'INSERT INTO buys_history (date, seller, article, qt, not_payed_yet, pay_date) VALUES("{}", "{}", "{}", {}, {}, "{}") ;'
                    .format(today, self.buy_choose_seller_comboBox.currentText(), self.buy_choose_article_comboBox.currentText(),
                            self.buy_article_qt.value()
                            , self.buy_article_price.value() * self.buy_article_qt.value(),
                            str(self.buy_date_to_pay.date().toPyDate())))
            else:

                curs.execute('''
                UPDATE articles SET qt = qt - {} WHERE name LIKE "{}"
                '''.format(self.buy_article_qt.value(), self.buy_choose_article_comboBox.currentText()))

                curs.execute(
                    'UPDATE sellers SET total_recived = total_recived + {}, pay_date = "{}"  WHERE F_name LIKE "{}" AND L_name LIKE "{}";'.format(
                        self.buy_article_price.value() * self.buy_article_qt.value(),
                        today
                        , self.buy_choose_seller_comboBox.currentText().split(' ')[0],
                        self.buy_choose_seller_comboBox.currentText().split(' ')[1]))

                curs.execute(
                    '''INSERT INTO buys_history
                     (date, seller, article, qt, payed, pay_date) 
                     VALUES("{}", "{}", "{}", {}, {}, "{}") ;
                     '''.format(today, self.buy_choose_seller_comboBox.currentText(),
                                self.buy_choose_article_comboBox.currentText(),
                                self.buy_article_qt.value()
                                , self.buy_article_price.value() * self.buy_article_qt.value(),
                                today))

            conn.commit()

    def open_sell_history(self):
        self.selles_history_wn = Selles_history()
        self.selles_history_wn.show()
        self.close()
        print('sell win opened')

    def open_buy_history(self):
        self.close()
        self.buy_history_wn = Buyes_history()
        self.buy_history_wn.show()

    def add_new_client(self):
        self.close()
        self.add_client_wn = ADD_new_client()
        self.add_client_wn.show()

    def edit_client(self):
        self.close()
        self.edite = Edite_client()
        self.edite.show()

    def add_new_seller(self):
        self.close()
        self.add_seller_wn = ADD_new_seller()
        self.add_seller_wn.show()

    def remove_seller(self):
        self.close()
        self.delete_seller_win = Remove_seller()
        self.delete_seller_win.show()

    def edit_seller(self):
        self.close()
        self.edite = Edite_seller()
        self.edite.show()

    def add_new_article(self):
        self.close()
        self.add_art_wn = ADD_new_art()
        self.add_art_wn.show()

    def remove_article(self):
        pass

    def edit_article(self):
        self.close()
        self.edit_art = Edite_art()
        self.edit_art.show()
    
    def settt(self):
        self.close()
        self.setti = Setting_wn()
        self.setti.show()

class Reset_pass(QMainWindow, first_open_win_dir):
    def __init__(self, parent=None):
        super(Reset_pass, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('إستعادة كلمة المرور ')
        self.save_btn.clicked.connect(self.reset__)
        self.safe_mod_checkBox.hide()
        # .setEchoMode(QtGui.QLineEdit.Password)
        self.pass_entry.setEchoMode(QLineEdit.Password)
        self.pass_confirmation.setEchoMode(QLineEdit.Password)

    def reset__(self):
        curs.execute('SELECT NAME FROM user;')
        curent_user_name = curs.fetchone()[0]
        print('the user name is : ', curent_user_name)

        if self.user_name_entry.text() == '':
            self.err = QErrorMessage()
            self.err.showMessage('يجب ملأ إسم المستخدم')
            self.err.setWindowTitle('خطأ في إسم المستخدم ')
            self.user_name_entry.setFocus()

        elif self.pass_entry.text() == '':
            self.err = QErrorMessage()
            self.err.showMessage('يجب ملأ كلمة مرور')
            self.err.setWindowTitle('خطأ في كلمة مرور')
            self.pass_entry.setFocus()

        elif self.pass_confirmation.text() == '':
            self.err = QErrorMessage()
            self.err.showMessage('يجب إعادة كلمة مرور')
            self.err.setWindowTitle('خطأ في إعادة كلمة مرور')
            self.pass_confirmation.setFocus()

        elif self.pass_entry.text() != self.pass_confirmation.text():
            self.err = QErrorMessage()
            self.err.showMessage('كلمات المرور غير متطابقة')
            self.err.setWindowTitle('خطأ في كلمة مرور')
            self.pass_confirmation.setFocus()

        elif self.user_name_entry.text() != curent_user_name:
            self.err = QErrorMessage()
            self.err.showMessage('لقد أدخلت إسم المستخدم خاطئ')
            self.err.setWindowTitle('إسم المستخدم خاطئ')
            self.user_name_entry.setFocus()

        else:
            curs.execute('SELECT PASSWORD FROM user ;')
            if self.pass_confirmation.text() == curs.fetchone()[0]:
                self.err = QErrorMessage()
                self.err.showMessage('يجب إختيار كلمة مرور لا تشبه القديمة')
                self.err.setWindowTitle('كلمة مرور موجودة بالفعل ')
                self.pass_entry.setFocus()
            else:
                curs.execute('UPDATE user SET PASSWORD = "{}";'.format(self.pass_confirmation.text()))
                conn.commit()
                self.back_to_app = VirificationAlert()
                self.back_to_app.show()
                self.close()

class VirificationAlert(QWidget, virification_alert_win_dir):
    def __init__(self, parent = None):
        super(VirificationAlert, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('تأكيد الهوية')
        self.setFixedSize(289,161)
        self.virif_btn.clicked.connect(self.virife_password)
        self.virif_pass_line.setEchoMode(QLineEdit.Password)
        self.help.mousePressEvent = self.get_help
        self.reset_password.mousePressEvent = self.get_reset_password

    def virife_password(self):
        curs.execute('SELECT PASSWORD FROM user')
        curent_pass = curs.fetchone()
        print(curent_pass[0])

        if self.virif_pass_line.text() !='':
            if curent_pass[0] == self.virif_pass_line.text():
                print('yes {} is same '.format(curent_pass[0]))
                self.home_wn = Home()
                self.home_wn.show()
                self.close()

            else:
                self.err = QtWidgets.QErrorMessage()
                self.err.showMessage('كلمة مرور غير صحيحة ')
                self.err.setWindowTitle('خطأ')
                self.virif_pass_line.setFocus()
                self.virif_pass_line.setText('')
        else:
            self.err = QtWidgets.QErrorMessage()
            self.err.showMessage('كلمة مرور غير صالحة ')
            self.err.setWindowTitle('خطأ')
            self.virif_pass_line.setFocus()

    def get_reset_password(self, event):
        print('the reset password was clicked')
        self.reset_pass_wn = Reset_pass()
        self.reset_pass_wn.show()
        self.close()



    def get_help(self, event):
        print('the help was clicked')

class Setting_wn(QWidget, setting_win_dir):
    def __init__(self, parent = None):
        super(Setting_wn, self).__init__(parent)
        QWidget.__init__(self)
        self.setupUi(self)
        self.setWindowFlag(QtCore.Qt.WindowCloseButtonHint, False)
        self.setWindowTitle('إعدادات')
        self.bbb.setSource(QtCore.QUrl.fromLocalFile("src/ttt.html"))
        self.current_pass.textChanged.connect(self.onTextChange)
        self.confirm_pass.textChanged.connect(self.onTChange)
        self.save_pass.setEnabled(False)
        self.back_btn.clicked.connect(self.back__)
        self.save_pass.clicked.connect(self.save_thePass)
        curs.execute('SELECT safe FROM tools')
        self.safe = curs.fetchone()[0]
        if self.safe == 1 :
            self.current_pass.setEnabled(False)
            self.new_pass.setEnabled(True)
            self.confirm_pass.setEnabled(True)
        
        self.ref()
        self.save_user_info.clicked.connect(self.save_info)

    
    
    def save_info(self):
        curs.execute('UPDATE user SET NAME = "{}", market_name = "{}", phone = "{}", adress = "{}"'.format(self.user_name.text(), self.market_name.text(), 
        str(self.phone.text()), self.adress.text()))
        conn.commit()
        self.ref()

    def ref(self):
        curs.execute('SELECT NAME, market_name, phone, adress FROM user')
        info = []
        for i in curs.fetchone():
            info.append(i)
    
        self.user_name.setText(info[0])
        self.market_name.setText(info[1])
        self.phone.setText(info[2])
        self.adress.setText(info[3])
            


    def save_thePass(self):
        if self.safe_check.isChecked():
            curs.execute('UPDATE tools SET safe = 1')
            curs.execute('UPDATE user SET PASSWORD = NULL')
        else:
            curs.execute('UPDATE user SET PASSWORD = "{}"'.format(self.confirm_pass.text()))
            curs.execute('UPDATE tools SET safe = 0')
        conn.commit()
        self.current_pass.setText('')
        self.new_pass.setText('')
        self.confirm_pass.setText('')

    def back__(self):
        self.close()
        self.home = Home()
        self.home.show()

    def onTextChange(self):
        if self.current_pass.text() == '' or self.current_pass.text() == ' ':
            self.save_pass.setEnabled(False)

        else:
            curs.execute('SELECT PASSWORD FROM user')
            if curs.fetchone()[0] == self.current_pass.text():
                self.new_pass.setEnabled(True)
                self.confirm_pass.setEnabled(True)
                self.safe_check.setEnabled(True)
            else:
                self.new_pass.setEnabled(False)
                self.confirm_pass.setEnabled(False)
                self.safe_check.setEnabled(False)

    def onTChange(self):
        if self.confirm_pass.text() == self.new_pass.text():
            self.save_pass.setEnabled(True)
        else:
            self.save_pass.setEnabled(False)


def delay_counter():
    curs.execute('SELECT delay FROM tools;')
    delay = curs.fetchone()[0]
    while True:
        delay += 1
        time.sleep(1)
        print(delay)
        if delay == 180000:
            break

def main():
    app = QtWidgets.QApplication(sys.argv)
    splash_wn = Splash()
    splash_wn.show()
    sys.exit(app.exec_())
#
# pr1 = mp.Process(target=main)
# pr2 =mp.Process(target=delay_counter)
# pr1.start()
# pr2.start()
#///////////////////////////////////////////////////////////////////////////////////////////////////////////////////////

if __name__ == '__main__' :
    main()
